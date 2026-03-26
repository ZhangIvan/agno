import asyncio
import os
from pathlib import Path
from typing import IO, Any, List, Optional, Union
from uuid import uuid4

from agno.knowledge.chunking.document import DocumentChunking
from agno.knowledge.chunking.strategy import ChunkingStrategy, ChunkingStrategyType
from agno.knowledge.document.base import Document
from agno.knowledge.reader.base import Reader
from agno.knowledge.types import ContentType
from agno.utils.log import log_debug, log_error, log_warning

try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:
    raise ImportError("The `python-docx` package is not installed. Please install it via `pip install python-docx`.")


class DocxReader(Reader):
    """Reader for Doc/Docx files"""

    def __init__(
        self,
        chunking_strategy: Optional[ChunkingStrategy] = DocumentChunking(),
        capture_pages: bool = True,
        pages_cache_dir: Optional[str] = None,
        image_dpi: int = 150,
        **kwargs,
    ):
        self.capture_pages = capture_pages
        self.pages_cache_dir = pages_cache_dir or os.getenv("AGNO_PAGE_CACHE_DIR")

        self.image_dpi = image_dpi
        super().__init__(chunking_strategy=chunking_strategy, **kwargs)

    @classmethod
    def get_supported_chunking_strategies(cls) -> List[ChunkingStrategyType]:
        """Get the list of supported chunking strategies for DOCX readers."""
        return [
            ChunkingStrategyType.DOCUMENT_CHUNKER,
            ChunkingStrategyType.CODE_CHUNKER,
            ChunkingStrategyType.FIXED_SIZE_CHUNKER,
            ChunkingStrategyType.SEMANTIC_CHUNKER,
            ChunkingStrategyType.AGENTIC_CHUNKER,
            ChunkingStrategyType.RECURSIVE_CHUNKER,
        ]

    @classmethod
    def get_supported_content_types(cls) -> List[ContentType]:
        return [ContentType.DOCX, ContentType.DOC]

    def read(self, file: Union[Path, IO[Any]], name: Optional[str] = None) -> List[Document]:
        """Read a docx file and return a list of documents"""
        _tmp_capture_path: Optional[str] = None
        try:
            file_path: Optional[str] = None
            if isinstance(file, Path):
                if not file.exists():
                    raise FileNotFoundError(f"Could not find file: {file}")
                log_debug(f"Reading: {file}")
                docx_document = DocxDocument(str(file))
                doc_name = name or file.stem
                file_path = str(file)
            else:
                log_debug(f"Reading uploaded file: {getattr(file, 'name', 'BytesIO')}")
                docx_document = DocxDocument(file)
                doc_name = name or getattr(file, "name", "docx_file").split(".")[0]
                if self.capture_pages:
                    import os, shutil, tempfile
                    if hasattr(file, "seek"):
                        file.seek(0)
                    _tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                    shutil.copyfileobj(file, _tmp)
                    _tmp.close()
                    if hasattr(file, "seek"):
                        file.seek(0)
                    file_path = _tmp.name
                    _tmp_capture_path = _tmp.name

            doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

            documents: List[Document] = (
                [Document(name=doc_name, id=str(uuid4()), content=doc_content)] if doc_content else []
            )
            if self.chunk:
                chunked_documents: List[Document] = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                result = chunked_documents
            else:
                result = documents

            # Capture page images if requested
            if self.capture_pages and file_path and result:
                try:
                    from agno.knowledge.reader.page_capture import capture_docx_pages, get_page_cache_dir
                    cache_dir = get_page_cache_dir(self.pages_cache_dir, doc_name)
                    page_images = capture_docx_pages(file_path, cache_dir, dpi=self.image_dpi)
                    total_pages = len(page_images)

                    # Assign page images to chunks by proportional mapping
                    total_chunks = len(result)
                    for chunk_index, doc in enumerate(result):
                        if total_pages > 0 and total_chunks > 0:
                            # Map chunk index → approximate page number (1-based)
                            approx_page = max(1, round((chunk_index / total_chunks) * total_pages) + 1)
                            approx_page = min(approx_page, total_pages)
                            doc.meta_data["page_number"] = approx_page
                            doc.meta_data["total_pages"] = total_pages
                            doc.meta_data["doc_type"] = "text_chunk"
                            if approx_page in page_images:
                                doc.meta_data["page_image_path"] = page_images[approx_page]
                                doc.meta_data["pages_cache_dir"] = cache_dir

                    # Append image documents for multimodal embedding
                    for page_num, image_path in page_images.items():
                        result.append(
                            Document(
                                name=doc_name,
                                id=f"{doc_name}_img_{page_num}",
                                content="",
                                meta_data={
                                    "doc_type": "page_image",
                                    "page_number": page_num,
                                    "total_pages": total_pages,
                                    "page_image_path": image_path,
                                    "pages_cache_dir": cache_dir,
                                },
                                content_id=f"{doc_name}_page_{page_num}",
                            )
                        )
                except Exception as e:
                    log_warning(f"Failed to capture DOCX page images: {e}")

            return result

        except Exception as e:
            log_error(f"Error reading file: {e}")
            raise ValueError(f"Error reading file: {e}")
        finally:
            if _tmp_capture_path:
                try:
                    import os
                    os.unlink(_tmp_capture_path)
                except OSError:
                    pass

    async def async_read(self, file: Union[Path, IO[Any]], name: Optional[str] = None) -> List[Document]:
        """Asynchronously read a docx file and return a list of documents"""
        try:
            return await asyncio.to_thread(self.read, file, name)
        except Exception as e:
            log_error(f"Error reading file asynchronously: {e}")
            raise ValueError(f"Error reading file asynchronously: {e}")
